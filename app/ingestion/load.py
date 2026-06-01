"""Parse a .docx user manual into a flat, ordered element stream.

The output is consumed by `app/ingestion/chunk.py` (which groups elements
into chunks) and `app/ingestion/images.py` (which writes image blobs to
disk). This module is purely structural — it does not produce chunks
itself and does not write any files.

Why python-docx and not Docling:
Our manuals are clean .docx with proper Word heading styles. Docling's
PDF-layout parser adds ~500 MB of ML dependencies that we wouldn't use.
python-docx is ~5 MB and gives us exactly what we need: paragraph styles,
run-level formatting (bold/italic), and inline image references.
"""
from __future__ import annotations

import logging
import re
from dataclasses import dataclass, field
from pathlib import Path
from typing import Literal

from docx import Document as DocxDocument
from docx.document import Document
from docx.text.paragraph import Paragraph

log = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Element model
# ---------------------------------------------------------------------------

ElementKind = Literal[
    "heading",          # Heading 1/2/3 paragraph
    "paragraph",        # Normal body paragraph
    "list_item",        # List Paragraph style (numbered or bulleted)
    "image",            # An inline image (no associated text yet)
    "figure_caption",   # A paragraph whose text starts with "Figure:"
    "callout",          # A paragraph that starts with "Note:" or "Caution:"
    "when_to_use",      # The italic "When to use:" line under a procedure heading
]


@dataclass
class Element:
    kind: ElementKind
    text: str = ""

    # Heading level (1, 2, 3); 0 for non-headings.
    level: int = 0

    # Text of every run in this paragraph that is bold. Used by the chunker
    # to build the chunk's ui_labels field.
    bold_runs: list[str] = field(default_factory=list)

    # For callouts: "note" or "caution".
    callout_type: str = ""

    # For image elements: the raw binary + extension. images.py writes
    # these to disk; the path is filled in afterwards.
    image_blob: bytes | None = None
    image_extension: str = ""

    # Original 0-based paragraph index in the source document. Useful for
    # ordering images against their following figure captions.
    raw_index: int = 0


@dataclass
class LoadedDocument:
    elements: list[Element]
    source_path: Path
    doc_title: str

    # Slug derived from doc_title, used for the per-document image folder
    # (e.g. "patient-management").
    doc_slug: str


# ---------------------------------------------------------------------------
# Title-block detection
# ---------------------------------------------------------------------------
# The first 3 non-empty paragraphs of every manual are a title block:
#   Patient Management Module
#   EMR System — User Manual for Health Workers
#   Ministry of Health · Uganda
# We skip them — they aren't content, they're branding. The patterns are
# loose enough to match across modules (Doctor Module, Lab Module, etc.).
TITLE_BLOCK_PATTERNS = [
    re.compile(r"^[A-Z][A-Za-z ]+ Module$"),
    re.compile(r"^EMR System\s*[—–-]\s*User Manual"),
    re.compile(r"^Ministry of Health"),
]


def _is_title_block_line(text: str) -> bool:
    for pat in TITLE_BLOCK_PATTERNS:
        if pat.match(text.strip()):
            return True
    return False


# ---------------------------------------------------------------------------
# Paragraph classification
# ---------------------------------------------------------------------------

_FIGURE_PREFIX = re.compile(r"^\s*Figure\s*:", re.IGNORECASE)
_NOTE_PREFIX = re.compile(r"^\s*Note\s*:", re.IGNORECASE)
_CAUTION_PREFIX = re.compile(r"^\s*Caution\s*:", re.IGNORECASE)
_WHEN_TO_USE_PREFIX = re.compile(r"^\s*When\s+to\s+use\s*:", re.IGNORECASE)


def _classify_paragraph(p: Paragraph) -> ElementKind | None:
    """Return the element kind for this paragraph, or None to skip it
    (empty paragraphs and the like)."""
    text = p.text.strip()
    if not text:
        return None

    style_name = p.style.name if p.style is not None else ""

    if style_name.startswith("Heading"):
        return "heading"
    if _FIGURE_PREFIX.match(text):
        return "figure_caption"
    if _NOTE_PREFIX.match(text) or _CAUTION_PREFIX.match(text):
        return "callout"
    if _WHEN_TO_USE_PREFIX.match(text):
        return "when_to_use"
    if style_name == "List Paragraph":
        return "list_item"
    return "paragraph"


def _heading_level(p: Paragraph) -> int:
    """Extract the heading level from a 'Heading N' style name. Returns 0 if not a heading."""
    style_name = p.style.name if p.style is not None else ""
    if style_name.startswith("Heading"):
        try:
            return int(style_name.split()[-1])
        except (ValueError, IndexError):
            return 0
    return 0


def _bold_runs(p: Paragraph) -> list[str]:
    """Return the text of every bold run in this paragraph."""
    runs = []
    for r in p.runs:
        if r.bold and r.text.strip():
            runs.append(r.text.strip())
    return runs


def _callout_type(text: str) -> str:
    if _CAUTION_PREFIX.match(text):
        return "caution"
    if _NOTE_PREFIX.match(text):
        return "note"
    return ""


# ---------------------------------------------------------------------------
# Image extraction (from within paragraphs)
# ---------------------------------------------------------------------------
# python-docx exposes inline images via the document's part relationships.
# Each <w:drawing><wp:inline> in a paragraph corresponds to one image.
# We iterate the paragraph's runs and pull image blob bytes via the relationship.

def _images_in_paragraph(p: Paragraph, doc: Document) -> list[tuple[bytes, str]]:
    """Return [(image_bytes, extension), ...] for every inline image in this paragraph."""
    out: list[tuple[bytes, str]] = []
    # The XML namespace used by Word for drawings.
    # blip = "Binary Large Image / Picture"
    NS = {
        "w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main",
        "a": "http://schemas.openxmlformats.org/drawingml/2006/main",
        "r": "http://schemas.openxmlformats.org/officeDocument/2006/relationships",
        "wp": "http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing",
    }
    # Find every <a:blip r:embed="rIdXX"/> inside this paragraph's XML.
    for blip in p._p.iter("{http://schemas.openxmlformats.org/drawingml/2006/main}blip"):
        embed_id = blip.get(
            "{http://schemas.openxmlformats.org/officeDocument/2006/relationships}embed"
        )
        if not embed_id:
            continue
        rel = p.part.rels.get(embed_id)
        if rel is None or "image" not in rel.reltype.lower():
            continue
        image_part = rel.target_part
        blob = image_part.blob
        ext = image_part.partname.split(".")[-1].lower() if "." in image_part.partname else "png"
        out.append((blob, ext))
    return out


# ---------------------------------------------------------------------------
# Public entry point
# ---------------------------------------------------------------------------

def _slugify(title: str) -> str:
    """Lower-case, replace non-alphanum with hyphens, collapse repeats."""
    s = re.sub(r"[^A-Za-z0-9]+", "-", title.lower()).strip("-")
    return s or "untitled"


def _extract_doc_title(paragraphs: list[Paragraph]) -> str:
    """The doc title is the first non-empty paragraph (typically the title block's first line).
    Used for the image folder slug and for the documents.title column."""
    for p in paragraphs:
        text = p.text.strip()
        if text:
            return text
    return "untitled"


def load_docx(path: Path) -> LoadedDocument:
    """Parse a .docx file into a flat, ordered list of Elements.

    Side effects: none. This function does not touch the filesystem outside
    of reading `path`.

    The title-block (first 3 paragraphs matching the title-block pattern)
    is skipped. Page headers/footers — stored in Word's section.header /
    section.footer regions, not in the body — are simply not iterated, so
    they're skipped automatically.
    """
    log.info("Loading docx: %s", path)
    doc: Document = DocxDocument(str(path))
    paragraphs = list(doc.paragraphs)

    doc_title = _extract_doc_title(paragraphs)
    doc_slug = _slugify(doc_title)

    elements: list[Element] = []
    title_block_skipped = 0

    for idx, p in enumerate(paragraphs):
        text = p.text.strip()

        # Drop the title-block lines (up to 3 of them) from the front.
        if title_block_skipped < 3 and text and _is_title_block_line(text):
            log.debug("Skipping title-block paragraph: %r", text)
            title_block_skipped += 1
            continue

        # Emit any inline images in this paragraph as their own elements,
        # in document order, BEFORE the paragraph's text element.
        # This way, when a paragraph contains both an image and trailing
        # text, the image appears in the stream before subsequent captions.
        for blob, ext in _images_in_paragraph(p, doc):
            elements.append(Element(
                kind="image",
                image_blob=blob,
                image_extension=ext,
                raw_index=idx,
            ))

        kind = _classify_paragraph(p)
        if kind is None:
            continue

        elem = Element(
            kind=kind,
            text=text,
            level=_heading_level(p) if kind == "heading" else 0,
            bold_runs=_bold_runs(p),
            callout_type=_callout_type(text) if kind == "callout" else "",
            raw_index=idx,
        )
        elements.append(elem)

    log.info(
        "Loaded %s: %d elements (%d images, %d figure_captions, %d headings)",
        path.name,
        len(elements),
        sum(1 for e in elements if e.kind == "image"),
        sum(1 for e in elements if e.kind == "figure_caption"),
        sum(1 for e in elements if e.kind == "heading"),
    )

    return LoadedDocument(
        elements=elements,
        source_path=path,
        doc_title=doc_title,
        doc_slug=doc_slug,
    )
